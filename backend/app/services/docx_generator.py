import copy
import zipfile
from pathlib import Path

from lxml import etree


W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
NS = {"w": W_NS}


FIELD_ALIASES = {
    "project.name": "project_name",
    "document.code": "report_no",
    "reportHeader.reportNo": "report_no",
    "reportHeader.customer": "customer",
    "reportHeader.sample": "sample",
    "reportHeader.conclusion": "conclusion",
}


def _replace_content_controls(xml_bytes: bytes, values: dict[str, str]) -> bytes:
    parser = etree.XMLParser(remove_blank_text=False)
    root = etree.fromstring(xml_bytes, parser)
    for control in root.xpath(".//w:sdt", namespaces=NS):
        tags = control.xpath("./w:sdtPr/w:tag/@w:val", namespaces=NS)
        if not tags or tags[0] not in values:
            continue
        texts = control.xpath("./w:sdtContent//w:t", namespaces=NS)
        if not texts:
            continue
        texts[0].text = values[tags[0]] or ""
        for text in texts[1:]:
            text.text = ""
    return etree.tostring(root, xml_declaration=True, encoding="UTF-8", standalone=True)


def _create_simple_document(output: Path, data: dict) -> None:
    template = Path(__file__).with_name("_blank.docx")
    del template
    from docx import Document

    document = Document()
    document.add_heading(data.get("project_name") or "分析报告", 0)
    table = document.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    for label, key in (("报告编号", "report_no"), ("客户名称", "customer"), ("样品名称", "sample")):
        cells = table.add_row().cells
        cells[0].text = label
        cells[1].text = data.get(key, "")
    document.add_heading("结论", level=1)
    document.add_paragraph(data.get("conclusion") or "")
    document.save(output)


def generate_docx(template_path: Path, output: Path, data: dict) -> None:
    output.parent.mkdir(parents=True, exist_ok=True)
    if not template_path.exists():
        _create_simple_document(output, data)
        return

    values = {tag: str(data.get(field) or "") for tag, field in FIELD_ALIASES.items()}
    with zipfile.ZipFile(template_path, "r") as source, zipfile.ZipFile(output, "w", zipfile.ZIP_DEFLATED) as target:
        for item in source.infolist():
            content = source.read(item.filename)
            normalized_name = item.filename.replace("\\", "/")
            if normalized_name == "word/document.xml" or (
                normalized_name.startswith("word/header") and normalized_name.endswith(".xml")
            ):
                content = _replace_content_controls(content, values)
            normalized_info = copy.copy(item)
            normalized_info.filename = normalized_name
            target.writestr(normalized_info, content)
    _append_test_items(output, data.get("test_items", []))


def _append_test_items(output: Path, items: list[dict]) -> None:
    if not items:
        return
    from docx import Document
    from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT

    document = Document(output)
    document.add_heading("检测结果", level=1)
    table = document.add_table(rows=1, cols=6)
    table.style = "Table Grid"
    for cell, label in zip(table.rows[0].cells, ("分类", "检测项目", "检测方法", "技术要求", "结果", "结论")):
        cell.text = label
    category_cells: dict[str, list] = {}
    for item in items:
        cells = table.add_row().cells
        values = (
            item.get("category", ""), item.get("name", ""), item.get("method", ""),
            item.get("requirement", ""), f"{item.get('result', '')} {item.get('unit', '')}".strip(),
            item.get("conclusion", ""),
        )
        for cell, value in zip(cells, values):
            cell.text = str(value)
        category_cells.setdefault(str(item.get("category", "")), []).append(cells[0])
    for cells in category_cells.values():
        if len(cells) > 1:
            merged = cells[0]
            for cell in cells[1:]:
                merged = merged.merge(cell)
            merged.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    document.save(output)
